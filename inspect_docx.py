import docx

doc = docx.Document('필기양식.docx')
for i, p in enumerate(doc.paragraphs):
    print(f"Paragraph {i}: {p.text}")

for i, table in enumerate(doc.tables):
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            print(f"Table {i}, Row {r}, Cell {c}: {cell.text}")
